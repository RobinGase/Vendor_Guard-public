import json
from pathlib import Path
from docx import Document
from agents.base import invoke_chat_model
from models.finding import Finding, VendorProfile

MODEL = "claude-opus-4-6"

MEMO_PROMPT = """You are a senior IT audit manager writing a formal vendor risk assessment memo.

Vendor: {vendor_name} | Sector: {sector} | Services: {services}

RAG Scorecard:
{scorecard}

Findings summary:
{findings_summary}

Write a professional audit memo with these sections:
1. Executive Summary (3-5 sentences: overall risk verdict, most critical issues, overall recommendation)
2. Assessment Scope (frameworks assessed, what documents were reviewed)
3. Key Findings (narrative of the most significant gaps, grouped by framework)
4. Recommendations (numbered list, prioritized by severity)
5. Conclusion

Use formal audit language. Reference specific control IDs and framework names. Do not use em-dashes.
"""


def _findings_summary(findings: list[Finding]) -> str:
    lines = []
    for f in findings:
        lines.append(f"[{f.framework}] {f.control_id} - {f.status} ({f.severity}): {f.evidence}")
    return "\n".join(lines)


def _scorecard_summary(scorecard: dict) -> str:
    lines = []
    for fw, data in scorecard.items():
        lines.append(f"{fw}: {data['rag']} (gaps: {data.get('gaps', 0)}, critical: {data.get('critical', 0)})")
    return "\n".join(lines)


def draft_memo_text(
    profile: VendorProfile,
    findings: list[Finding],
    scorecard: dict,
) -> str:
    """Call Claude to draft the audit memo narrative. Returns plain text."""
    prompt = MEMO_PROMPT.format(
        vendor_name=profile.name,
        sector=profile.sector,
        services=", ".join(profile.services),
        scorecard=_scorecard_summary(scorecard),
        findings_summary=_findings_summary(findings),
    )
    return invoke_chat_model(
        model=MODEL,
        max_tokens=2048,
        user_prompt=prompt,
    )


def _fallback_memo_text(
    profile: VendorProfile,
    findings: list[Finding],
    scorecard: dict,
) -> str:
    red_frameworks = [fw for fw, data in scorecard.items() if data.get("rag") == "Red"]
    amber_frameworks = [fw for fw, data in scorecard.items() if data.get("rag") == "Amber"]
    top_findings = findings[:5]

    lines = [
        "Executive Summary",
        f"Vendor {profile.name} was assessed across {len(scorecard)} framework(s).",
        f"Overall high-risk frameworks: {', '.join(red_frameworks) if red_frameworks else 'None'}.",
        f"Frameworks requiring follow-up: {', '.join(amber_frameworks) if amber_frameworks else 'None'}.",
        "",
        "Assessment Scope",
        f"Sector: {profile.sector}.",
        f"Services: {', '.join(profile.services)}.",
        "",
        "Key Findings",
    ]

    if top_findings:
        for finding in top_findings:
            lines.append(
                f"- [{finding.framework}] {finding.control_id}: {finding.status} ({finding.severity}). {finding.evidence}"
            )
    else:
        lines.append("- No findings were produced.")

    lines.extend([
        "",
        "Recommendations",
        "1. Review the highlighted gaps and partial controls.",
        "2. Validate the underlying evidence with the vendor.",
        "3. Track remediation actions for all high-risk findings.",
        "",
        "Conclusion",
        "This memo was generated using the built-in fallback summary because AI narrative drafting was unavailable during this run.",
    ])
    return "\n".join(lines)


def write_audit_memo(
    profile: VendorProfile,
    findings: list[Finding],
    scorecard: dict,
    output_path: Path,
    memo_text: str = None,
) -> str:
    """Write audit memo as DOCX. Returns the memo text for reuse."""
    if memo_text is None:
        try:
            memo_text = draft_memo_text(profile, findings, scorecard)
        except Exception:
            memo_text = _fallback_memo_text(profile, findings, scorecard)

    doc = Document()
    doc.add_heading(f"Vendor Risk Assessment: {profile.name}", 0)
    doc.add_heading("Audit Memo", level=1)

    for line in memo_text.split("\n"):
        if line.strip():
            doc.add_paragraph(line)
        else:
            doc.add_paragraph("")

    doc.save(str(output_path))
    return memo_text
